#!/usr/bin/env python
# -*- coding: utf-8 -*-

# 将 ALGORITHM_DESIGN_SPEC.md 转换为 Word 文档（.docx）
# 依赖：python-docx
# 使用（Windows PowerShell）：
#   pip install python-docx; python .\scripts\generate_docx.py

import os
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD_PATH = os.path.join(ROOT, 'ALGORITHM_DESIGN_SPEC.md')
DOCX_PATH = os.path.join(ROOT, 'ALGORITHM_DESIGN_SPEC.docx')


def add_code_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    font = run.font
    font.name = 'Courier New'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
    font.size = Pt(10.5)


def markdown_to_docx(md_text: str, doc: Document):
    lines = md_text.splitlines()
    in_code = False
    code_buffer = []

    bullet_re = re.compile(r"^\s*[-*]\s+")
    number_re = re.compile(r"^\s*\d+\.\s+")

    for line in lines:
        # 处理代码块围栏
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_buffer = []
            else:
                # 结束代码块
                in_code = False
                add_code_paragraph(doc, "\n".join(code_buffer))
            continue

        if in_code:
            code_buffer.append(line)
            continue

        # 标题
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            continue

        # 无序列表
        if bullet_re.match(line):
            doc.add_paragraph(bullet_re.sub('', line), style='List Bullet')
            continue

        # 有序列表
        if number_re.match(line):
            doc.add_paragraph(number_re.sub('', line), style='List Number')
            continue

        # 普通段落
        doc.add_paragraph(line)


def main():
    if not os.path.exists(MD_PATH):
        raise SystemExit(f"找不到Markdown文件: {MD_PATH}")
    with open(MD_PATH, 'r', encoding='utf-8') as f:
        md = f.read()
    doc = Document()
    markdown_to_docx(md, doc)
    doc.save(DOCX_PATH)
    print(f"已生成: {DOCX_PATH}")


if __name__ == '__main__':
    main()
